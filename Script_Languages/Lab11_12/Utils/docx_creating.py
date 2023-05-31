import os
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt
from matplotlib import pyplot as plt

from Utils.book_scrapping import graph_paragraph, count_words


def add_basic_info(book, doc, report_author):
    p = doc.add_paragraph()

    p.add_run(f"Title: ").font.size = Pt(20)
    run = p.add_run(f"{book.title}\n")
    run.font.size = Pt(20)
    run.font.bold = True

    p.add_run(f"Book Author: ").font.size = Pt(20)
    run = p.add_run(f"{book.author}\n")
    run.font.size = Pt(20)
    run.font.bold = True

    p.add_run(f"Report Author: ").font.size = Pt(20)
    run = p.add_run(f"{report_author}\n")
    run.font.size = Pt(20)
    run.font.bold = True


def add_picture(doc, picture_path, width=10.0):
    p = doc.add_paragraph()

    p.add_run().add_picture(picture_path, width=Cm(width))
    p.alignment = 1


def add_charts(book, doc):
    temp_picture_path = "output-files/temp_fig.png"

    for i, chapter in enumerate(book.chapters):
        fig = graph_paragraph(book.paragraphs_lengths[i], i)
        fig.savefig(temp_picture_path, bbox_inches="tight")

        p = doc.add_paragraph()
        p.alignment = 1
        chapter_roman = arabic_to_roman(i + 1)
        run = p.add_run(f"Chapter {chapter_roman}\n\n")
        run.font.size = Pt(25)
        run.bold = True

        add_picture(doc, temp_picture_path, 15.9)
        doc.add_page_break()

    os.remove(temp_picture_path)


def add_final_page(book, doc):
    plt.figure(figsize=(8, 4))

    plt.title(f"Number of Paragraphs in Chapters")
    plt.xlabel("Chapter")
    plt.ylabel("Paragraphs")
    plt.xticks(range(1, len(book.paragraphs_lengths) + 1, 2))

    chapters_lengths = [len(x) for x in book.paragraphs_lengths]
    plt.bar(range(1, len(book.paragraphs_lengths) + 1), chapters_lengths)

    temp_picture_path = "output-files/temp_fig.png"
    plt.savefig(temp_picture_path, bbox_inches="tight")
    plt.close()

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f"Summary Page\n\n")
    run.font.size = Pt(25)
    run.font.bold = True

    add_picture(doc, temp_picture_path, 15.9)
    os.remove(temp_picture_path)

    headers = ["Min Words in Paragraph", "Max Words in Paragraph", "Total Number of Paragraphs",
               "Avg. Number of Paragraphs in Chapters", "Avg. Number of Words in Paragraphs",
               "Total number of words", "Avg. Number of Words in Chapters"]
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    values = []

    exact_paragraph_lengths = count_words(book.chapters, round_down=False)
    min_chapter = min([x[0] for x in exact_paragraph_lengths])
    max_chapter = max([x[-1] for x in exact_paragraph_lengths])

    no_of_paragraphs = sum([len(chapter) for chapter in exact_paragraph_lengths])
    avg_no_of_paragraphs_in_chapters = no_of_paragraphs / len(book.paragraphs_lengths)

    sum_of_lengths_of_paragraphs = sum([sum(chapter) for chapter in exact_paragraph_lengths])
    avg_length_of_paragraphs = sum_of_lengths_of_paragraphs / no_of_paragraphs

    no_of_words = sum(sum(chapter) for chapter in exact_paragraph_lengths)
    avg_length_of_chapters = no_of_words / len(book.paragraphs_lengths)

    values.append(min_chapter)
    values.append(max_chapter)
    values.append(no_of_paragraphs)
    values.append(int(avg_no_of_paragraphs_in_chapters))
    values.append(int(avg_length_of_paragraphs))
    values.append(no_of_words)
    values.append(int(avg_length_of_chapters))

    for i, cell in enumerate(values):
        cell = table.cell(1, i)
        cell.text = str(values[i])


def arabic_to_roman(number):
    arabic_to_roman_mapping = {
        1000: 'M',
        900: 'CM',
        500: 'D',
        400: 'CD',
        100: 'C',
        90: 'XC',
        50: 'L',
        40: 'XL',
        10: 'X',
        9: 'IX',
        5: 'V',
        4: 'IV',
        1: 'I'
    }

    roman_number = ""
    for arabic, roman in arabic_to_roman_mapping.items():
        while number >= arabic:
            roman_number += roman
            number -= arabic

    return roman_number


def create_docx(book, cfg):
    report_author = cfg.params.author
    image_path = cfg.picture.path[:-4]+".png"
    doc = Document()

    add_basic_info(book, doc, report_author)

    if os.path.exists(image_path):
        add_picture(doc, image_path, 10)

    doc.add_page_break()

    add_charts(book, doc)

    add_final_page(book, doc)

    formatted_datetime = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc.save(f"{cfg.docx.path}report_{formatted_datetime}_{book.id}.docx")


